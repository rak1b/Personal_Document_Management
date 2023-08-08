import PyPDF2
from docx import Document
from reportlab.lib.pagesizes import letter, landscape
from reportlab.lib.units import inch
from reportlab.pdfgen import canvas

def convert_pdf_to_docx(pdf_path, docx_path):
    pdf_text = ""
    with open(pdf_path, "rb") as pdf_file:
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        num_pages = len(pdf_reader.pages)
        for page_num in range(num_pages):
            page = pdf_reader.pages[page_num]
            pdf_text += page.extract_text()
    doc = Document()
    doc.add_paragraph(pdf_text)
    doc.save(docx_path)


def convert_docx_to_pdf(docx_path, pdf_path):
    doc = Document(docx_path)
    pdf_canvas = canvas.Canvas(pdf_path, pagesize=landscape(letter))

    for para in doc.paragraphs:
        pdf_canvas.drawString(1 * inch, (10 - 0.2) * inch, para.text)
        pdf_canvas.showPage()
    pdf_canvas.save()
    print("convert 959945")
